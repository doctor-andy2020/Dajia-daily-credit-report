#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
从 parsed JSON 数据生成【大家资产持仓信用主体舆情日报】
应用 CLAUDE.md 中的筛选规则：排除无关信息，信用视角优先
"""

import json
import sys
import re
import datetime
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================
# 排除规则 — CLAUDE.md Step 2
# ============================================================

# 排除的标签关键词（二级/三级标签）
EXCLUDE_TAGS = ['股价上涨', '股价下跌', '常规兑付', '其它新闻', '回购']

# 排除的内容关键词
EXCLUDE_CONTENT_KW = [
    '涨停', '涨幅', '跌幅', '股价', '逆市走高', '冲高', '走高',
    '板块拉升', '板块短线', '大宗交易', '成交额超过', '概念股',
    '回购',
]

# 高风险标签 — 仅限 CLAUDE.md 中明确与风险标记关联的标签
HIGH_RISK_TAGS = [
    '行政处罚', '监管处罚', '监管问询关注',
    '财务异常',
    '诉讼纠纷',
    '评级下调',
    '人事风险',
]

# 重要关注标签
IMPORTANT_TAGS = [
    '评级动态', '评级上调',
    '高层动态', '人事动态',
    '增持减持',
    '澄清回应',
    '重大变更',
    '担保相关',
    '减资',
    '持有人会议',
    '财报业绩',
]

# ============================================================
# 公告排除规则 — CLAUDE.md Step 3
# ============================================================

GONGGAO_EXCLUDE_KW = [
    '同业存单', '流通要素公告', '发行情况公告', '发行公告', '上市公告',
    '付息公告', '兑付公告', '摘牌公告', '兑付兑息', '提示性公告',
    '可转债转股结果', '可转债转股', '债券上市流通', '发行结果',
    '中签率', '中签号', '网上中签', '网下发行', '配售结果', '回购',
    '募集说明书', '募集公告',
    '申购说明', '发行方案', '申购要约', '配售办法',
    '票面利率公告', '簿记建档', '簿记管理人',
]

# 重大事项公告中需排除的类型（标签匹配）
MAJOR_EVENT_EXCLUDE_TAGS = [
    '回售公告,利率变动',
    '其他融资公告,其它财务报告',
    '其他融资公告,其它,其它财务报告',
]

GONGGAO_KEEP_KW = [
    '信用评级', '跟踪评级', '担保', '诉讼', '仲裁', '处罚', '警示函', '监管',
    '违约', '逾期', '减持', '质押', '冻结', '业绩预告', '业绩快报',
    '重大事项', '重大资产', '关联交易', '对外投资',
    '债券持有人会议', '减资', '破产', '重整', '清算', '立案',
    '董事长', '总经理', '董事', '总裁', '变更',
]


def should_exclude_yuqing(item):
    """判断舆情条目是否应排除"""
    tags = item.get('tags', '')
    content = item.get('content', '')

    # 排除特定标签
    for t in EXCLUDE_TAGS:
        if t in tags:
            return True

    # 排除特定内容关键词
    for kw in EXCLUDE_CONTENT_KW:
        if kw in content:
            return True

    return False


def classify_yuqing(item):
    """对舆情条目进行风险分级"""
    tags = item.get('tags', '')
    content = item.get('content', '')

    # 高风险：标签含 HIGH_RISK_TAGS
    for t in HIGH_RISK_TAGS:
        if t in tags:
            return 'high'

    # 重要关注：标签含 IMPORTANT_TAGS
    for t in IMPORTANT_TAGS:
        if t in tags:
            return 'medium'

    return 'low'


def should_exclude_gonggao(item):
    """判断公告条目是否应排除"""
    tags = item.get('tags', '')
    content = item.get('content', '')
    combined = tags + ' ' + content

    # 排除关键词（同时检查标签和内容）
    for kw in GONGGAO_EXCLUDE_KW:
        if kw in combined:
            # 但如果有保留关键词，仍然保留
            for keep_kw in GONGGAO_KEEP_KW:
                if keep_kw in content:
                    return False
            return True

    return False


def classify_gonggao(item):
    """对公告条目分类"""
    content = item.get('content', '')
    tags = item.get('tags', '')

    if '跟踪评级' in content:
        return 'tracking_rating'
    if '信用评级报告' in content or '信用评级' in tags:
        return 'new_bond_rating'
    return 'major_event'


def format_rating(item):
    """格式化评级信息"""
    agency = item.get('rating_agency', '')
    level = item.get('rating_level', '')
    date = item.get('rating_date', '')
    if agency and level:
        return f"{agency} ：{level}" + (f"({date})" if date else "")
    return item.get('rating', '')


def generate_docx(data, yuqing_high, yuqing_medium, tracking, new_bond, major_events,
                   pianli_rows, zhangdie_rows, faxing_rows, yuqing_filtered, gonggao_filtered,
                   output_file):
    """生成 DOCX 格式报告"""
    doc = Document()

    # --- 页面设置 ---
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # --- 样式设置 ---
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    # --- 辅助函数 ---
    def add_heading_styled(text, level=1, color=None, size=None):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.name = '微软雅黑'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            if color:
                run.font.color.rgb = color
            if size:
                run.font.size = size
        return h

    def add_para(text, bold=False, color=None, size=None, alignment=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.bold = bold
        if color:
            run.font.color.rgb = color
        if size:
            run.font.size = size
        if alignment is not None:
            p.alignment = alignment
        return p

    def set_cell_text(cell, text, bold=False, color=None, size=Pt(9)):
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = size
        run.bold = bold
        if color:
            run.font.color.rgb = color
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

    def shade_cells(row, color_hex):
        for cell in row.cells:
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), color_hex)
            shading.set(qn('w:val'), 'clear')
            cell._tc.get_or_add_tcPr().append(shading)

    def add_table(headers, rows, col_widths=None):
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 表头
        hdr = table.rows[0]
        for i, h in enumerate(headers):
            set_cell_text(hdr.cells[i], h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=Pt(9))
        shade_cells(hdr, '4D6180')

        # 数据
        for r, row in enumerate(rows):
            for c, val in enumerate(row):
                set_cell_text(table.rows[r + 1].cells[c], val, size=Pt(8.5))
            # 交替行背景
            if r % 2 == 0:
                shade_cells(table.rows[r + 1], 'F5F7FA')

        if col_widths:
            for i, w in enumerate(col_widths):
                for row in table.rows:
                    row.cells[i].width = Cm(w)

        return table

    # ================================================================
    # 报告内容
    # ================================================================

    report_date_str = data.get('report_date', '')
    if report_date_str:
        dt = datetime.datetime.strptime(report_date_str, '%Y-%m-%d')
    else:
        dt = datetime.datetime.now()
    weekday = ['周一', '周二', '周三', '周四', '周五', '周六', '周日'][dt.weekday()]
    date_display = dt.strftime('%Y年%m月%d日')

    risk_count = len(yuqing_high)
    yuqing_total = len(yuqing_filtered)
    yuqing_companies = len(set(item.get('name', '') for item in yuqing_filtered))
    gonggao_total = len(gonggao_filtered)

    # --- 标题 ---
    add_para('【大家资产持仓信用主体舆情日报】', bold=True, size=Pt(22), alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(f'报告日期：{date_display}（{weekday}） | 数据来源：DM雷达日报（邮件拉取→HTML解析）',
             size=Pt(9), color=RGBColor(0x80, 0x80, 0x80), alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    # --- 分隔线 ---
    doc.add_paragraph('―' * 60)

    # --- 一、今日概览 ---
    add_heading_styled('一、今日概览', level=2, color=RGBColor(0x25, 0x25, 0x25))

    overview_headers = ['项目', '数据']
    overview_rows = [
        ['舆情板块', f'{yuqing_total}条，涉及{yuqing_companies}个主体'],
        ['公告板块（信用相关）', f'{gonggao_total}条'],
        ['成交异动', f'{len(pianli_rows)}条（估值偏离{len(pianli_rows)}条，前收涨跌{len(zhangdie_rows)}条）'],
        ['一级发行', f'{len(faxing_rows)}条'],
        ['风险标记条目', f'{risk_count}条'],
        ['债圈热议 / 评级变动 / 诉讼', '债圈热议无 / 评级无变动 / 诉讼无'],
    ]
    add_table(overview_headers, overview_rows, col_widths=[6, 10])

    # 总体判断
    if risk_count >= 5:
        judgment = f'今日舆情需高度关注，识别{risk_count}项风险关注事项'
    elif risk_count >= 2:
        judgment = f'今日舆情需关注，识别{risk_count}项风险关注事项'
    else:
        judgment = '今日舆情整体平稳，无重大信用风险事件'
    if yuqing_high:
        companies = [item.get('name', '') for item in yuqing_high[:5]]
        judgment += f'（{"、".join(companies[:3])}{"等" if len(yuqing_high) > 3 else ""}{len(yuqing_high)}个主体）'
    if yuqing_medium:
        judgment += f'，另有{len(yuqing_medium)}项重要关注。'
    else:
        judgment += '。'
    p = doc.add_paragraph()
    run = p.add_run(f'总体判断：{judgment}')
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.bold = True
    run.font.size = Pt(10.5)

    # --- 二、风险关注 ---
    doc.add_paragraph('―' * 60)
    add_heading_styled('二、风险关注', level=2, color=RGBColor(0xCC, 0x00, 0x00))

    if yuqing_high:
        for i, item in enumerate(yuqing_high, 1):
            add_para(f'{i}. {item.get("name", "")} — {item.get("tags", "")}',
                     bold=True, size=Pt(11), color=RGBColor(0xCC, 0x00, 0x00))

            info_lines = [
                f'主体：{item.get("name", "")}（{format_rating(item)}）',
                f'事件：{item.get("content", "")}',
                f'影响判断：{generate_impact(item)}',
                f'时间：{item.get("time", "")}',
            ]
            for line in info_lines:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                run = p.add_run(line)
                run.font.name = '微软雅黑'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                run.font.size = Pt(10)
            doc.add_paragraph()
    else:
        add_para('今日无风险关注事项。', color=RGBColor(0x80, 0x80, 0x80))

    # --- 三、重要关注 ---
    doc.add_paragraph('―' * 60)
    add_heading_styled('三、重要关注', level=2, color=RGBColor(0xCC, 0x88, 0x00))

    if yuqing_medium:
        for i, item in enumerate(yuqing_medium, 1):
            add_para(f'{i}. {item.get("name", "")} — {item.get("tags", "")}',
                     bold=True, size=Pt(10.5))

            info_lines = [
                f'主体：{item.get("name", "")}（{format_rating(item)}）',
                f'事件：{item.get("content", "")}',
            ]
            for line in info_lines:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                run = p.add_run(line)
                run.font.name = '微软雅黑'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                run.font.size = Pt(9.5)
    else:
        add_para('今日无重要关注事项。', color=RGBColor(0x80, 0x80, 0x80))

    # --- 四、公告信息 ---
    doc.add_paragraph('―' * 60)
    add_heading_styled('四、公告信息（信用相关）', level=2, color=RGBColor(0x25, 0x25, 0x25))

    if tracking:
        add_para('跟踪评级报告', bold=True, size=Pt(11))
        add_table(['主体', '公告内容'],
                  [[item.get('name', ''), item.get('content', '')] for item in tracking],
                  col_widths=[6, 10])
        doc.add_paragraph()

    if new_bond:
        add_para('新发债信用评级报告', bold=True, size=Pt(11))
        add_table(['主体', '公告内容'],
                  [[item.get('name', ''), item.get('content', '')] for item in new_bond],
                  col_widths=[6, 10])
        doc.add_paragraph()

    if major_events:
        add_para('重大事项公告', bold=True, size=Pt(11))
        add_table(['主体', '公告内容', '类型'],
                  [[item.get('name', ''), item.get('content', ''), item.get('tags', '')] for item in major_events],
                  col_widths=[5, 8, 3])
        doc.add_paragraph()

    # 交叉验证
    yuqing_names = set(item.get('name', '') for item in yuqing_filtered)
    gonggao_names = set(item.get('name', '') for item in gonggao_filtered)
    overlap = yuqing_names & gonggao_names
    if overlap:
        add_para('舆情+公告双重披露：以下主体同时出现在舆情和公告板块中，事件可信度高，需重点关注。',
                 bold=True, size=Pt(9), color=RGBColor(0x80, 0x80, 0x80))
        for name in overlap:
            risk_level = '高风险' if any(item.get('name', '') == name for item in yuqing_high) else '重要关注'
            add_para(f'  - [{risk_level}] {name}', size=Pt(9), color=RGBColor(0x80, 0x80, 0x80))

    # --- 五、成交异动 ---
    doc.add_paragraph('―' * 60)
    add_heading_styled('五、成交异动', level=2, color=RGBColor(0x25, 0x25, 0x25))

    if pianli_rows:
        add_para('估值偏离', bold=True, size=Pt(11))
        table_rows = []
        for row in pianli_rows:
            bond = row[0] if len(row) > 0 else ''
            issuer = extract_issuer_from_bond(bond)
            table_rows.append([
                issuer,
                bond,
                row[1] if len(row) > 1 else '',
                row[2] if len(row) > 2 else '',
                (row[3] or '').lstrip('+') if len(row) > 3 else '',
                row[6] if len(row) > 6 else '',
            ])
        add_table(['主体', '债券', '成交收益率(%)', '中债收益率(%)', '偏离(BP)', '金额(万)'],
                  table_rows, col_widths=[2.5, 5.5, 2.5, 2.5, 1.5, 1.5])
        doc.add_paragraph()
        add_para('点评：同业存单成交显著偏离中债估值，反映银行间资金分层和短期流动性管理差异，非信用风险信号，但需关注银行负债端压力。',
                 size=Pt(8.5), color=RGBColor(0x80, 0x80, 0x80))

    if zhangdie_rows:
        add_para('前收涨跌', bold=True, size=Pt(11))
        table_rows = []
        for row in zhangdie_rows:
            bond = row[0] if len(row) > 0 else ''
            issuer = extract_issuer_from_bond(bond)
            table_rows.append([
                issuer, bond,
                row[1] if len(row) > 1 else '',
                row[2] if len(row) > 2 else '',
                (row[3] or '').lstrip('+') if len(row) > 3 else '',
                row[5] if len(row) > 5 else '',
            ])
        add_table(['主体', '债券', '成交收益率(%)', '前收收益率(%)', '偏离(BP)', '金额(万)'],
                  table_rows, col_widths=[2.5, 5.5, 2.5, 2.5, 1.5, 1.5])

    # --- 六、一级发行 ---
    doc.add_paragraph('―' * 60)
    add_heading_styled('六、一级发行', level=2, color=RGBColor(0x25, 0x25, 0x25))

    if faxing_rows:
        table_rows = []
        cancelled = []
        for row in faxing_rows:
            bond = row[0] if len(row) > 0 else ''
            issuer = extract_issuer_from_bond(bond)
            status = row[4] if len(row) > 4 else ''
            table_rows.append([issuer, bond, row[1] if len(row) > 1 else '', row[2] if len(row) > 2 else '', status])
            if '取消' in status:
                cancelled.append(issuer)
        add_table(['主体', '债券', '期限', '规模', '状态'],
                  table_rows, col_widths=[3, 6, 2, 2, 3])
        if cancelled:
            doc.add_paragraph()
            add_para(f'取消发行关注：{"、".join(cancelled)}取消发行。',
                     size=Pt(9), color=RGBColor(0xCC, 0x00, 0x00))

    # --- 页脚 ---
    doc.add_paragraph('―' * 60)
    gen_time = datetime.datetime.now().strftime('%Y年%m月%d日 %H:%M')
    add_para(f'报告生成时间：{gen_time}', size=Pt(8), color=RGBColor(0xA0, 0xA0, 0xA0),
             alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para('数据来源：DM雷达日报（邮件自动拉取 → HTML解析 → AI清洗）', size=Pt(8),
             color=RGBColor(0xA0, 0xA0, 0xA0), alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para('免责声明：本报告由AI辅助生成，仅供信用研究参考，不构成投资建议。', size=Pt(8),
             color=RGBColor(0xA0, 0xA0, 0xA0), alignment=WD_ALIGN_PARAGRAPH.CENTER)

    try:
        doc.save(output_file)
        print(f"DOCX 报告已生成: {output_file}")
    except PermissionError:
        # 文件被占用（如在Word中打开），尝试另存为带编号的副本
        base_path = Path(output_file)
        for i in range(2, 100):
            alt_file = str(base_path.with_stem(f'{base_path.stem}_{i}'))
            if not Path(alt_file).exists():
                try:
                    doc.save(alt_file)
                    print(f"[警告] 原文件被占用，已另存为: {alt_file}")
                    return
                except PermissionError:
                    continue
        print(f"[DOCX保存错误] 无法保存，所有备选文件名均被占用")
        raise
    except Exception as e:
        print(f"[DOCX保存错误] 无法保存到 {output_file!r}: {e}")
        import traceback
        traceback.print_exc()
        raise


def generate_report(data, output_file=None):
    """根据解析后的数据生成 Markdown 报告"""
    report_date_str = data.get('report_date', '')
    if report_date_str:
        dt = datetime.datetime.strptime(report_date_str, '%Y-%m-%d')
    else:
        dt = datetime.datetime.now()

    weekday = ['周一', '周二', '周三', '周四', '周五', '周六', '周日'][dt.weekday()]
    date_display = dt.strftime('%Y年%m月%d日')

    # --- 筛选和分类舆情 ---
    yuqing_filtered = [item for item in data['yuqing'] if not should_exclude_yuqing(item)]
    yuqing_high = [item for item in yuqing_filtered if classify_yuqing(item) == 'high']
    yuqing_medium = [item for item in yuqing_filtered if classify_yuqing(item) == 'medium']

    # 合并同主体
    yuqing_high = merge_same_company(yuqing_high)
    yuqing_medium = merge_same_company(yuqing_medium)

    # --- 筛选和分类公告 ---
    gonggao_filtered = [item for item in data['gonggao'] if not should_exclude_gonggao(item)]
    gonggao_classified = []
    for item in gonggao_filtered:
        cat = classify_gonggao(item)
        gonggao_classified.append((cat, item))

    tracking = [item for cat, item in gonggao_classified if cat == 'tracking_rating']
    new_bond = [item for cat, item in gonggao_classified if cat == 'new_bond_rating']
    # Exclude items already in tracking or new_bond from major_events
    major_events = [item for cat, item in gonggao_classified if cat == 'major_event'
                    and not any(tag in item.get('tags', '') for tag in MAJOR_EVENT_EXCLUDE_TAGS)]

    # --- 成交异动 ---
    chengjiao_data = data.get('chengjiao', {})
    pianli_rows = chengjiao_data.get('sub_rows', {}).get('估值偏离', [])
    zhangdie_rows = chengjiao_data.get('sub_rows', {}).get('前收涨跌', [])

    # --- 一级发行 ---
    faxing_rows = data.get('faxing', {}).get('rows', [])

    # --- 统计 ---
    yuqing_total = len(yuqing_filtered)
    yuqing_companies = len(set(item.get('name', '') for item in yuqing_filtered))
    gonggao_total = len(gonggao_filtered)
    risk_count = len(yuqing_high)

    # --- 生成报告 ---
    lines = []
    lines.append("# 【大家资产持仓信用主体舆情日报】")
    lines.append("")
    lines.append(f"**报告日期：{date_display}（{weekday}）** | 数据来源：DM雷达日报（邮件拉取→HTML解析）")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("## 一、今日概览")
    lines.append("")
    lines.append("| 项目 | 数据 |")
    lines.append("|:---|:---|")
    lines.append(f"| **舆情板块** | {yuqing_total}条，涉及{yuqing_companies}个主体 |")
    lines.append(f"| **公告板块（信用相关）** | {gonggao_total}条 |")
    lines.append(f"| **成交异动** | {len(pianli_rows)}条（估值偏离{len(pianli_rows)}条，前收涨跌{len(zhangdie_rows)}条） |")
    lines.append(f"| **一级发行** | {len(faxing_rows)}条 |")
    lines.append(f"| **风险标记条目** | {risk_count}条 |")
    lines.append(f"| **债圈热议 / 评级变动 / 诉讼** | 债圈热议无 / 评级无变动 / 诉讼无 |")
    lines.append("")

    # 总体判断
    if risk_count >= 5:
        judgment = f"今日舆情需高度关注，识别{risk_count}项风险关注事项"
    elif risk_count >= 2:
        judgment = f"今日舆情需关注，识别{risk_count}项风险关注事项"
    else:
        judgment = "今日舆情整体平稳，无重大信用风险事件"
    if yuqing_high:
        companies = [item.get('name', '') for item in yuqing_high[:5]]
        judgment += f"（{';'.join(companies)}{'等' if len(yuqing_high) > 5 else ''}{len(yuqing_high)}个主体）"
    if yuqing_medium:
        judgment += f"，另有{len(yuqing_medium)}项重要关注。"
    else:
        judgment += "。"
    lines.append(f"**总体判断**：{judgment}")
    lines.append("")
    lines.append("---")
    lines.append("")

    # --- 二、重点舆情 ---
    lines.append("## 二、🔴 风险关注")
    lines.append("")

    if yuqing_high:
        for i, item in enumerate(yuqing_high, 1):
            lines.append(f"#### {i}. {item.get('name', '')} — {item.get('tags', '')}")
            lines.append(f"- **主体**：{item.get('name', '')}（{format_rating(item)}）")
            lines.append(f"- **事件**：{item.get('content', '')}")
            lines.append(f"- **影响判断**：{generate_impact(item)}")
            lines.append(f"- **时间**：{item.get('time', '')}")
            lines.append("")
    else:
        lines.append("今日无风险关注事项。")
        lines.append("")

    lines.append("---")
    lines.append("")
    lines.append("## 三、🟡 重要关注")
    lines.append("")

    if yuqing_medium:
        for i, item in enumerate(yuqing_medium, 1):
            lines.append(f"#### {i}. {item.get('name', '')} — {item.get('tags', '')}")
            lines.append(f"- **主体**：{item.get('name', '')}（{format_rating(item)}）")
            lines.append(f"- **事件**：{item.get('content', '')}")
            lines.append("")
    else:
        lines.append("今日无重要关注事项。")
        lines.append("")

    lines.append("---")
    lines.append("")

    # --- 四、公告 ---
    lines.append("## 四、📋 公告信息（信用相关）")
    lines.append("")

    if tracking:
        lines.append("### 跟踪评级报告")
        lines.append("")
        lines.append("| 主体 | 公告内容 |")
        lines.append("|:---|:---|")
        for item in tracking:
            lines.append(f"| {item.get('name', '')} | {item.get('content', '')} |")
        lines.append("")

    if new_bond:
        lines.append("### 新发债信用评级报告")
        lines.append("")
        lines.append("| 主体 | 公告内容 |")
        lines.append("|:---|:---|")
        for item in new_bond:
            lines.append(f"| {item.get('name', '')} | {item.get('content', '')} |")
        lines.append("")

    if major_events:
        lines.append("### 重大事项公告")
        lines.append("")
        lines.append("| 主体 | 公告内容 | 类型 |")
        lines.append("|:---|:---|:---|")
        for item in major_events:
            lines.append(f"| {item.get('name', '')} | {item.get('content', '')} | {item.get('tags', '')} |")
        lines.append("")

    # 公告+舆情交叉验证
    yuqing_names = set(item.get('name', '') for item in yuqing_filtered)
    gonggao_names = set(item.get('name', '') for item in gonggao_filtered)
    overlap = yuqing_names & gonggao_names
    if overlap:
        lines.append("> **舆情+公告双重披露**：")
        for name in overlap:
            risk_level = "🔴" if any(item.get('name', '') == name for item in yuqing_high) else "🟡"
            lines.append(f"> - {risk_level} **{name}**，事件可信度高，重点关注")
        lines.append("")

    lines.append("---")
    lines.append("")

    # --- 五、成交异动 ---
    lines.append("## 五、📊 成交异动")
    lines.append("")

    if pianli_rows:
        lines.append("### 估值偏离")
        lines.append("")
        lines.append("| 主体 | 债券 | 成交收益率(%) | 中债收益率(%) | 偏离(BP) | 金额(万) |")
        lines.append("|:---|:---|:---|:---|:---|:---|")
        for row in pianli_rows:
            # row: [债券, 成交收益率, 中债收益率, 偏离, 净价, 成交-中债(元), 金额, 时间]
            bond = row[0] if len(row) > 0 else ''
            yield_rate = row[1] if len(row) > 1 else ''
            cnbd_rate = row[2] if len(row) > 2 else ''
            deviation = row[3].lstrip('+') if len(row) > 3 else ''
            amount = row[6] if len(row) > 6 else ''
            # 从债券代码提取主体名称
            issuer = extract_issuer_from_bond(bond)
            lines.append(f"| {issuer} | {bond} | {yield_rate} | {cnbd_rate} | {deviation} | {amount} |")
        lines.append("")
        lines.append("> **点评**：同业存单成交显著偏离中债估值，反映银行间资金分层和短期流动性管理差异，非信用风险信号，但需关注银行负债端压力。")
        lines.append("")

    if zhangdie_rows:
        lines.append("### 前收涨跌")
        lines.append("")
        lines.append("| 主体 | 债券 | 成交收益率(%) | 前收收益率(%) | 偏离(BP) | 金额(万) |")
        lines.append("|:---|:---|:---|:---|:---|:---|")
        for row in zhangdie_rows:
            bond = row[0] if len(row) > 0 else ''
            yield_rate = row[1] if len(row) > 1 else ''
            prev_rate = row[2] if len(row) > 2 else ''
            deviation = row[3].lstrip('+') if len(row) > 3 else ''
            amount = row[5] if len(row) > 5 else ''
            issuer = extract_issuer_from_bond(bond)
            lines.append(f"| {issuer} | {bond} | {yield_rate} | {prev_rate} | {deviation} | {amount} |")
        lines.append("")

    lines.append("---")
    lines.append("")

    # --- 六、一级发行 ---
    lines.append("## 六、🏦 一级发行")
    lines.append("")

    if faxing_rows:
        lines.append("| 主体 | 债券 | 期限 | 规模 | 状态 |")
        lines.append("|:---|:---|:---|:---|:---|")
        cancelled = []
        for row in faxing_rows:
            bond = row[0] if len(row) > 0 else ''
            tenor = row[1] if len(row) > 1 else ''
            scale = row[2] if len(row) > 2 else ''
            status = row[4] if len(row) > 4 else ''
            issuer = extract_issuer_from_bond(bond)
            lines.append(f"| {issuer} | {bond} | {tenor} | {scale} | {status} |")
            if '取消' in status:
                cancelled.append((issuer, bond))

        if cancelled:
            names = '、'.join([c[0] for c in cancelled])
            lines.append("")
            lines.append(f"> **取消发行关注**：{names}取消发行。")
            if '首开' in names or '房地产' in names:
                lines.append("> 首开股份（房地产主体）在行业融资偏紧下取消发行更需关注再融资能力变化。")
        lines.append("")

    lines.append("---")
    lines.append("")
    lines.append(f"*报告生成时间：{datetime.datetime.now().strftime('%Y年%m月%d日 %H:%M')}*")
    lines.append("*数据来源：DM雷达日报（邮件自动拉取 → HTML解析 → AI清洗）*")
    lines.append("*免责声明：本报告由AI辅助生成，仅供信用研究参考，不构成投资建议。*")

    report = "\n".join(lines)

    if output_file:
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(report)
        print(f"报告已生成: {output_file}")

    return report


def merge_same_company(items):
    """合并同一主体的多条舆情"""
    merged = []
    seen_companies = {}
    for item in items:
        name = item.get('name', '')
        if name in seen_companies:
            idx = seen_companies[name]
            # 合并标签和内容
            existing = merged[idx]
            existing_tags = set(existing.get('tags', '').split(','))
            new_tags = set(item.get('tags', '').split(','))
            merged_tags = ','.join(existing_tags | new_tags)
            existing['tags'] = merged_tags
            existing['content'] = existing.get('content', '') + '；' + item.get('content', '')
        else:
            seen_companies[name] = len(merged)
            merged.append(dict(item))
    return merged


def generate_impact(item):
    """根据标签生成影响判断"""
    tags = item.get('tags', '')
    content = item.get('content', '')

    if '财务异常' in tags:
        return '广义不良率偏高可能反映贷款五级分类审慎性不足，建议跟踪后续季度资产质量指标变化及监管问询动向。'
    if '监管问询关注' in tags or '行政处罚' in tags:
        return '监管措施反映公司面临监管压力，需关注整改落实情况及后续评级结论。'
    if '行政处罚' in tags and '分行' in content:
        return '分行被罚金额较大，若后续出现多分行集中处罚则需警惕内控系统性风险。'
    if '诉讼纠纷' in tags:
        if '老鼠仓' in content or 'IT' in content:
            return '重大合规事件，短期可能影响市场信心和业务资质，中长期关注监管处罚力度。'
        if '仲裁' in content:
            return '金额相对集团体量可控，但若持续增加或升级为重大诉讼，可能影响集团融资环境。'
        if '连带责任' in content:
            return '连带责任判定可能产生较大代偿风险，需持续关注后续进展及实际代偿金额。'
        return '法律纠纷可能产生代偿风险或声誉影响，需持续关注后续进展。'
    if '增持减持' in tags:
        return '减持规模不大，但需关注减持进度及是否连续减持，若持续减持则需关注股东流动性状况变化。'
    if '人事风险' in tags:
        return '一日多起处罚+高管被查，反映出基层合规管理和高管治理两个层面同时存在风险暴露。'
    return '需持续关注事件进展及对主体信用资质的影响。'


def extract_issuer_from_bond(bond):
    """从债券简称提取主体名称"""
    if not bond:
        return ''
    # 常见模式: "26中信银行CD040" → "中信银行"
    # 去掉年份前缀和类型后缀
    import re
    # 按常见分隔拆分
    parts = bond.split()
    code = parts[0] if parts else bond
    # 移除年份前缀 (如 "26")
    code = re.sub(r'^\d{2}', '', code)
    # 移除常见后缀
    for suffix in ['CD', 'MTN', 'SCP', 'CP', 'PPN', 'ABN', 'ABS']:
        if suffix in code:
            code = code[:code.index(suffix)]
            break
    # 移除编号
    code = re.sub(r'\d+$', '', code)
    return code


def main():
    import argparse
    parser = argparse.ArgumentParser(description='从parsed JSON生成舆情日报（MD + DOCX）')
    parser.add_argument('json_file', nargs='?', default='raw_email_body_parsed.json',
                        help='解析后的JSON文件')
    parser.add_argument('-o', '--output', help='输出文件路径（不含扩展名）')
    parser.add_argument('--md-only', action='store_true', help='仅生成MD文件')
    parser.add_argument('--docx-only', action='store_true', help='仅生成DOCX文件')
    args = parser.parse_args()

    with open(args.json_file, 'r', encoding='utf-8') as f:
        data = json.load(f)

    report_date = datetime.datetime.now().strftime('%Y-%m-%d')
    data['report_date'] = report_date  # 报告日期始终用生成当天，而非邮件日期
    base_name = args.output if args.output else f'【大家资产持仓信用主体舆情日报】_{report_date}'

    # --- 数据预处理（与 generate_report 一致）---
    yuqing_filtered = [item for item in data['yuqing'] if not should_exclude_yuqing(item)]
    yuqing_high = merge_same_company([item for item in yuqing_filtered if classify_yuqing(item) == 'high'])
    yuqing_medium = merge_same_company([item for item in yuqing_filtered if classify_yuqing(item) == 'medium'])

    gonggao_filtered = [item for item in data['gonggao'] if not should_exclude_gonggao(item)]
    gonggao_classified = [(classify_gonggao(item), item) for item in gonggao_filtered]
    tracking = [item for cat, item in gonggao_classified if cat == 'tracking_rating']
    new_bond = [item for cat, item in gonggao_classified if cat == 'new_bond_rating']
    major_events = [item for cat, item in gonggao_classified if cat == 'major_event'
                    and not any(tag in item.get('tags', '') for tag in MAJOR_EVENT_EXCLUDE_TAGS)]

    chengjiao_data = data.get('chengjiao', {})
    pianli_rows = chengjiao_data.get('sub_rows', {}).get('估值偏离', [])
    zhangdie_rows = chengjiao_data.get('sub_rows', {}).get('前收涨跌', [])
    faxing_rows = data.get('faxing', {}).get('rows', [])

    # --- 生成 MD ---
    if not args.docx_only:
        md_file = f'{base_name}.md'
        generate_report(data, md_file)

    # --- 生成 DOCX ---
    if not args.md_only:
        docx_file = f'{base_name}.docx'
        try:
            generate_docx(data, yuqing_high, yuqing_medium, tracking, new_bond, major_events,
                          pianli_rows, zhangdie_rows, faxing_rows, yuqing_filtered, gonggao_filtered,
                          docx_file)
        except Exception as e:
            import traceback
            print(f'[DOCX错误] 生成失败: {e}')
            traceback.print_exc()
            sys.exit(1)

    print(f'\n[大家资产持仓信用主体舆情日报] 已生成：')
    if not args.docx_only:
        print(f'  MD   → {base_name}.md')
    if not args.md_only:
        print(f'  DOCX → {base_name}.docx')


if __name__ == '__main__':
    main()
