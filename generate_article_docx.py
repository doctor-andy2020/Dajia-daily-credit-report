"""Generate a formatted DOCX from extracted DM article text.

Supports TWO article formats:
  1. 债市要闻速览 — sections use 【】 brackets (e.g. 【宏观要闻】)
  2. DM信用早报   — sections use plain-text headers (e.g. 宏观要闻)
                    with sub-sections (国内/地方/海外, 地产/城投/其它)
"""
import os, re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUT = os.path.dirname(os.path.abspath(__file__))
INPUT = os.path.join(OUT, "dm_article_output.txt")

# ── Article type detection ──────────────────────────────────────────
def detect_article_type(lines):
    """Return 'zaobao' for DM信用早报, 'yaowen' for 债市要闻速览."""
    for i in range(min(50, len(lines))):
        s = lines[i].strip()
        if "信用早报" in s and ("DM" in s or "早报" in s):
            return "zaobao"
        if "债市要闻速览" in s:
            return "yaowen"
    # Fallback: check for 【】 brackets
    for line in lines[:100]:
        if line.strip().startswith("【"):
            return "yaowen"
    # Default to zaobao (newer format)
    return "zaobao"

# ── Text cleaning ───────────────────────────────────────────────────
def read_and_clean(path):
    with open(path, encoding="utf-8") as f:
        lines = f.readlines()

    article_type = detect_article_type(lines)

    # Find article start
    start = 0
    for i, line in enumerate(lines):
        s = line.strip()
        if re.search(r"(债市要闻速览|DM信用早报|信用早报)\d{4}", s):
            start = i
            break

    # Find article end
    end = len(lines)
    if article_type == "yaowen":
        for i in range(len(lines) - 1, start, -1):
            if "涉及区域" in lines[i] or "更多债市要闻详情" in lines[i]:
                end = i
                break
    else:  # zaobao
        for i in range(len(lines) - 1, start, -1):
            if "免责声明" in lines[i]:
                end = i
                break

    content = lines[start:end]

    # Menu artifacts to remove (both formats)
    MENU_ITEMS = {
        "菜单", "首页", "固收综合屏", "债券违约", "舆情", "城投地图",
        "利率债市场观点", "债市复盘", "行业分类", "利差曲线", "多资产综合屏",
        "评级调整", "商票逾期", "DM专栏", "金融板块", "自选雷达", "实时利差",
        "AI研报", "标记颜色", "刷新", "固定标签页", "关闭其他标签页",
        "关闭右边标签页", "收藏", "字号", "小", "中", "大", "区域名称",
        "中华人民共和国",
    }

    # DM信用早报 header artifacts (these are UI elements, not content)
    ZAOBAO_HEADER_SKIP = {"收藏"}

    cleaned = []
    for line in content:
        s = line.strip()

        # Skip empty / whitespace-only
        if not s:
            cleaned.append("")
            continue

        # Skip isolated numbers (menu items / page numbers)
        if s.isdigit() and len(s) <= 3:
            continue

        # Skip known menu items
        if s in MENU_ITEMS:
            continue

        # Skip user name
        if re.match(r"^孟思锜\s*$", s):
            continue

        # Skip zaobao header artifacts
        if article_type == "zaobao":
            if s in ZAOBAO_HEADER_SKIP:
                continue
            # Keep "来源：" and "标签  " lines — parse_sections needs them
            if re.match(r"^字号\s*$", s):
                continue
            # Skip the "来源：DM AI舆情查看原链接" line
            if s.startswith("来源：") and "查看原链接" in s:
                continue

        # Skip .xlsx file links
        if s.endswith(".xlsx"):
            continue

        cleaned.append(line.rstrip())

    # Merge consecutive empty lines
    merged = []
    prev_empty = False
    for line in cleaned:
        empty = not line.strip() or line.strip() == " "
        if empty and prev_empty:
            continue
        merged.append(line)
        prev_empty = empty

    return merged, article_type


# ── Section parsing ─────────────────────────────────────────────────
# Known section headers for DM信用早报 (plain text, no brackets)
ZAOBAO_SECTIONS = [
    "经济数据", "宏观要闻", "公司新闻",
    "取消发行/终止审核", "财报业绩/企业拿地",
    "评级动态", "人事动态", "金融机构信息汇总",
]
ZAOBAO_SUBS = ["国内", "地方", "海外", "地产", "城投", "其它"]

# Content-start markers — patterns that reliably indicate a line is
# ELABORATION/CONTENT, NOT a new news headline.
# IMPORTANT: do NOT include bare entity names (中国央行, 北京, 上海, etc.)
# as those will absorb unrelated headlines starting with the same entity.
CONTENT_STARTERS = [
    "据", "根据", "数据显示", "数据显",
    "新华社", "央视新闻", "中新社", "中新网", "人民日报",
    "当地时间", "此前", "此前一天", "同日", "另据",
    "公告显示", "公告显", "募集说", "上述", "该", "此次", "本次",
    "目前", "截至", "其中", "具体", "近日", "此外", "同时", "另外",
    "综合新华社", "路透社",
    "周五", "周六", "周日", "周一", "周二", "周三", "周四",
    "以下简称", "截至发稿", "据中国",
    # Purpose-clause openers — common in policy/regulation content paragraphs
    "为贯彻", "为贯彻落实", "为规范", "为促进", "为落实",
    "为加强", "为推动", "为进一步", "为更好", "为切实",
    "为深入", "为加快", "为保障", "为支持", "为引导",
    # Reporting-continuation patterns
    "修订说明指出", "修订说明", "修订后的",
    # Financial analysis / commentary continuation starters
    "展望", "预计", "预测", "我们认为", "分析人士", "分析认为",
    "值得关注", "受此影响", "这意味着", "据了解", "据悉",
    "事实上", "实际上", "相比之下", "有分析", "业内",
    "市场人士", "未来", "随着", "近期", "今年以来",
    "从中长期", "下一阶段", "往后看", "往前看",
    "在此背景下", "在此基础", "由此来看",
    "相较于", "对比", "与之相比",
    "换句话", "换言之", "也就", "可以说",
    "注意到", "需要关注", "需关注", "应关注",
    "反观", "相对应", "另一方面", "与之相对",
    "尤其", "特别", "更为关键", "更值得",
    "具体来看", "具体而言", "展开来看",
    "进一步", "进一", "不仅", "不只",
    "从结构看", "从数据看", "分项来看", "分行来看",
    "按行业", "按地区", "按业务",
    "对应", "对应地", "相应地",
    # Company-specific content elaboration
    "公司表示", "公司称", "公司指出", "公司预计", "公司预测",
    "公司计划", "公司拟", "公司已", "公司将", "公司正在",
    "该公", "该行", "该机构",
    "今年以", "去年以", "上半年", "下半年",
    "一季度", "二季度", "三季度", "四季度",
    "报告期", "同期", "环比", "同比",
    "全年", "本月", "本周", "本季",
    "最新", "近期", "近日",
]

# Reporting-verb suffixes — if an entity name ends with one of these
# followed by ：or , the line is a content/reporting clause, not a title.
# e.g. "中国央行数据显示，" → content  (ends with 显示)
#      "中国央行：M2增长..."   → title   (entity + ：without reporting verb)
REPORTING_VERBS = [
    "显示", "公布", "发布", "报道", "称", "表示", "公告",
    "指出", "强调", "披露", "透露", "介绍", "回应", "公布",
    "宣布", "声明", "通知", "通报",
]

# Title-indicator words — if a line starts with one of these, it is almost
# certainly a new headline, not content continuation.
TITLE_STARTERS = [
    "穆迪：", "惠誉：", "标普：", "中诚信：", "联合资信：", "大公：",
    "穆迪上调", "穆迪下调", "惠誉上调", "惠誉下调", "标普上调", "标普下调",
    "中诚信上调", "中诚信下调",
]


def _find_colon(text, max_pos=30):
    """Return position of first ：or : in text[:max_pos], or -1 if none."""
    for i, ch in enumerate(text):
        if i >= max_pos:
            break
        if ch in ("：", ":"):
            return i
    return -1


def _has_reporting_verb_before_colon(text, colon_pos):
    """Check if entity-name before colon ends with a reporting verb."""
    before = text[:colon_pos]
    return any(before.endswith(v) for v in REPORTING_VERBS)


def _same_topic(text, title, min_common=8):
    """Check if `text` elaborates on the same news topic as `title`.

    Both text and title are assumed to start with the same entity prefix
    (e.g. both start with "中国央行：").  Returns True if they share enough
    initial content to be considered the same news item.
    """
    # Compare the bodies (after ：/：)
    t_colon = _find_colon(title, 25)
    c_colon = _find_colon(text, 25)
    if t_colon < 0 or c_colon < 0:
        return False
    title_body = title[t_colon + 1:].strip()
    text_body = text[c_colon + 1:].strip()
    if not title_body or not text_body:
        return False
    # Count matching chars at start of body
    check_len = min(len(title_body), len(text_body), min_common)
    if check_len < 4:
        return False
    matches = sum(1 for i in range(check_len)
                  if title_body[i] == text_body[i])
    return matches >= check_len * 0.6  # ≥60% char overlap → same topic


def is_content_line(text, title_text=""):
    """Check if text looks like a content/follow-up paragraph, not a new title.

    Key heuristic:
    - "EntityName：statement" (entity + ：without reporting verb) → TITLE
      UNLESS it shares the same topic as title_text (then it's CONTENT).
    - "EntityName数据显示，" (entity + reporting verb + ，) → CONTENT
    - Lines starting with 据/根据/新华社 etc. → CONTENT
    """
    text = text.strip()
    if not text:
        return False

    # ── Strong TITLE indicators (check first — these are NEVER content) ──
    # 1. Starts with rating agency + ：pattern
    for ts in TITLE_STARTERS:
        if text.startswith(ts):
            return False

    # ── CONTENT indicators ──
    # Pattern 1: starts with known content/elaboration marker
    for marker in CONTENT_STARTERS:
        if text.startswith(marker):
            return True

    # Pattern 1b: starts with a date — content elaboration
    # e.g. "6月11日，欧洲央行宣布..."
    if re.match(r'^\d{1,2}月\d{1,2}日', text):
        return True

    # Pattern 2: contains reporting-verb + continuation in first 55 chars
    # e.g. "中国央行数据显示，" / "中国人民银行发布关于..."
    # The verb may be followed by punctuation (：:,，), whitespace, or
    # common clause-starters (关于, 对, 为, 将, 拟, 的, 了).
    m = re.search(r'(显示|公布|发布|报道|称|表示|公告|指出|强调|披露|透露|介绍|回应)[：:,，\s关于对为将拟的了]', text[:55])
    if m:
        entity_part = text[:m.start()]
        if len(entity_part) < 20:
            return True

    # Pattern 2b: shares a 《...》 document reference with the title
    # e.g. Title: "中基协修订发布《基金经理兼任...工作指引》"
    #      Text:  "为贯彻落实...对《基金经理兼任...工作指引》进行修订..."
    if title_text:
        title_refs = re.findall(r'《([^》]+)》', title_text)
        if title_refs:
            for ref in title_refs:
                if len(ref) >= 6 and ref in text:
                    return True

    # Pattern 2c: shares key topic phrase with title (≥6 common chars)
    # e.g. Title: "石澳道14号豪宅以5.63亿港元易手"
    #      Text:  "香港石澳一处历史悠久的洋房以5.63亿港元售出..."
    # IMPORTANT: if the title uses 《...》 to name a document, ONLY Pattern 2b
    # (exact doc name match) can confirm continuation — fuzzy phrase matching
    # could match a DIFFERENT document with a similar name.
    _title_has_book = '《' in title_text if title_text else False
    if title_text and len(title_text) >= 8 and not _title_has_book:
        for window in range(min(15, len(title_text)), 6, -1):
            for start in range(0, len(title_text) - window + 1):
                chunk = title_text[start:start + window]
                if chunk not in text:
                    continue
                if chunk[0] in '，。、：；（）':
                    continue
                if '《' in chunk or '》' in chunk:
                    continue
                tpos = text.index(chunk)
                # Skip chunks at position 0 in both if too long — shared
                # opening that could be same entity but different docs.
                if start == 0 and tpos == 0 and window > 10:
                    continue
                return True

    # ── Same-topic continuation (BEFORE ：title check) ──
    # Pattern 3: shares entity+topic with the current title.
    # "中国央行：M2增长8.6%。M1增长5.5%。" is content of
    # "中国央行：M2余额353.67万亿元" — same entity, same topic.
    if title_text and len(title_text) >= 3:
        t_colon = _find_colon(title_text, 25)
        c_colon = _find_colon(text, 25)
        # Both have ：at headline position with the same entity?
        if (t_colon > 0 and c_colon > 0
                and t_colon < 25 and c_colon < 25):
            title_entity = title_text[:t_colon]
            text_entity = text[:c_colon]
            if (len(title_entity) >= 2
                    and title_entity == text_entity
                    and not _has_reporting_verb_before_colon(text, c_colon)):
                # Same entity, both have ：→ check if same topic
                if _same_topic(text, title_text, min_common=8):
                    return True

    # ── TITLE-indicator: ：without reporting verb ──
    # After same-topic check above, a line with ：in headline position
    # that is NOT about the same topic is a NEW title.
    colon = _find_colon(text, 30)
    if colon > 0 and colon < 25:
        if not _has_reporting_verb_before_colon(text, colon):
            return False

    # ── Remaining prefix matching ──
    # Pattern 4: shares entity prefix with title (no ：in either)
    #
    # IMPORTANT: when the title contains 《...》, prefix matching is
    # suppressed — different docs from the same entity can share the
    # same opening (e.g. "中基协发布《基金A》" vs "中基协发布《基金B》").
    # Only Pattern 2b (exact doc name) can confirm same-topic for 《 titles.
    _title_has_book = '《' in title_text if title_text else False
    if title_text and len(title_text) >= 3 and not _title_has_book:
        t_colon2 = _find_colon(title_text, 25)
        if t_colon2 > 0:
            title_entity = title_text[:t_colon2]
            if len(title_entity) >= 2 and text.startswith(title_entity):
                if len(text) > len(title_text) * 0.7:
                    return True
        else:
            # Title has no colon — try prefix match with descending length.
            for n in [8, 6, 4]:
                if len(title_text) >= n:
                    prefix = title_text[:n]
                    if text.startswith(prefix):
                        return True
                else:
                    break
            # Last resort: 3-char entity prefix for very short titles,
            # but only if text is clearly longer (elaboration, not headline)
            if len(title_text) >= 3 and len(text) > len(title_text) * 1.3:
                if text.startswith(title_text[:3]):
                    return True

    # ── Fallback: long text following a title is almost certainly content ──
    # Real news titles in DM reports are rarely longer than 60 chars.
    # Long lines that follow an existing title are paragraphs, not new titles.
    if title_text and len(text) > 80:
        return True

    return False


def parse_sections(lines, article_type):
    """Parse article into structured sections.

    Returns dict with: title, date, tags, type, sections[]
    Each section: {heading, sub_heading, items: [{title, content}]}
    """
    article = {
        "title": "",
        "date": "",
        "tags": "",
        "type": article_type,
        "sections": [],
    }

    i = 0
    # ── Extract title ──
    while i < len(lines):
        s = lines[i].strip()
        if s and ("要闻速览" in s or "信用早报" in s or "早报" in s):
            # Truncate at | — keep only "DM信用早报0615" (zaobao)
            # or "债市要闻速览0612" (yaowen). The post-pipe content
            # is a summary of the day's headlines, not the title.
            if "|" in s:
                article["title"] = s.split("|")[0].strip()
            else:
                article["title"] = s
            i += 1
            break
        i += 1

    # ── Extract date ──
    if i < len(lines) and re.match(r"\d{4}-\d{2}-\d{2}", lines[i].strip()):
        article["date"] = lines[i].strip()
        i += 1

    # ── Skip header meta (来源, 收藏, 标签, 字号, etc.) ──
    # Only look ahead a limited number of lines to find the tag.
    SKIP_META = {
        "收藏", "来源：DM AI舆情查看原链接",
    }
    _tag_scan = 0
    while i < len(lines) and _tag_scan < 15:
        s = lines[i].strip()
        if s.startswith("标签") and "标签" in s:
            article["tags"] = s.replace("标签", "").strip()
            i += 1
            break
        if s in SKIP_META or s.startswith("来源："):
            i += 1
            _tag_scan += 1
            continue
        # Stop scanning if we hit a section header or non-meta content
        if s in ZAOBAO_SECTIONS or (article_type == "yaowen" and s.startswith("【")):
            break
        if s and not s.startswith("字号") and not re.match(r"^(小|中|大)$", s):
            # Non-meta content — stop scanning
            break
        i += 1
        _tag_scan += 1

    # ── Skip to first content section ──
    while i < len(lines):
        s = lines[i].strip()
        if article_type == "yaowen":
            if s.startswith("【"):
                break
        else:  # zaobao
            if s in ZAOBAO_SECTIONS:
                break
        i += 1

    # ── Parse sections ──
    current_section = None  # {heading, sub_heading, items: [{title, content}]}
    current_sub = None
    current_items = []
    current_title = None
    current_content_lines = []

    def flush_item():
        nonlocal current_title, current_content_lines
        if current_title:
            content_text = "\n".join(current_content_lines).strip()
            current_items.append({
                "title": current_title,
                "content": content_text,
            })
        current_title = None
        current_content_lines = []

    def flush_section():
        nonlocal current_section, current_sub, current_items
        flush_item()
        if current_items and current_section is not None:
            current_section["items"].extend(current_items)
        current_items = []
        current_sub = None

    while i < len(lines):
        s = lines[i].strip()

        # ── Section header detection ──
        is_section = False
        is_sub = False

        if article_type == "yaowen":
            is_section = s.startswith("【") and s.endswith("】")
        else:  # zaobao
            is_section = s in ZAOBAO_SECTIONS
            is_sub = s in ZAOBAO_SUBS

        if is_section:
            # Start new section
            flush_section()
            current_section = {"heading": s, "sub_heading": None, "items": []}
            article["sections"].append(current_section)
            current_sub = None
            i += 1
            # Skip blank lines after section header
            while i < len(lines) and not lines[i].strip():
                i += 1
            continue

        # ── Sub-section header (zaobao only) ──
        if is_sub:
            flush_section()  # Save accumulated items to previous sub-section
            # New sub-section under same parent
            parent_heading = current_section["heading"] if current_section else ""
            current_section = {"heading": parent_heading,
                               "sub_heading": s, "items": []}
            article["sections"].append(current_section)
            current_sub = s
            i += 1
            # Skip blank lines
            while i < len(lines) and not lines[i].strip():
                i += 1
            continue

        # ── NBSP line (u00a0) = item separator in zaobao, skip in both ──
        if s == " ":
            if current_title:
                flush_item()
            i += 1
            continue

        # ── Empty line ──
        if not s:
            if article_type == "zaobao":
                # In zaobao, empty lines between title and content are NOT boundaries.
                # Only &nbsp; separates items. Just skip and continue.
                i += 1
                continue
            else:
                # In yaowen, empty line = item boundary
                if current_title and current_content_lines:
                    flush_item()
                i += 1
                continue

        # ── Check for Gangtise delimiter (zaobao) ──
        if article_type == "zaobao" and "以下内容来自" in s and "Gangtise" in s:
            flush_item()
            flush_section()
            # Start a new section for 机构观点
            current_section = {"heading": "机构观点", "sub_heading": None, "items": []}
            article["sections"].append(current_section)
            current_sub = None
            i += 1
            continue

        # ── Content line: either new title or content ──
        if current_title and not current_content_lines:
            # We have a title; check if this line is content continuation
            if is_content_line(s, current_title):
                current_content_lines.append(s)
            else:
                # New title → flush previous (title-only OK for zaobao)
                current_items.append({
                    "title": current_title,
                    "content": "",
                })
                current_title = s
            i += 1
            continue

        if current_title and current_content_lines:
            # Already in content; more content or new title?
            if is_content_line(s, current_title):
                current_content_lines.append(s)
            else:
                flush_item()
                current_title = s
            i += 1
            continue

        # No current title yet → this line is a new title
        if not current_title:
            current_title = s
            i += 1
            continue

        i += 1

    # Flush remaining
    flush_item()
    flush_section()

    # Clean up empty sections (parent sections that only have sub-sections)
    article["sections"] = [s for s in article["sections"] if s["items"]]

    return article


# ── DOCX generation ─────────────────────────────────────────────────
def _add_heading_run(paragraph, text, size, bold=False, color=None, font_name="微软雅黑"):
    """Add a single run to a paragraph with formatting."""
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = font_name
    if color:
        run.font.color.rgb = color
    # Set East Asian font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn("w:rFonts"), {})
    rFonts.set(qn("w:eastAsia"), font_name)
    rPr.insert(0, rFonts)
    return run


def _add_body_paragraph(doc, text, indent=True, size=10, bold=False):
    """Add a body paragraph with optional first-line indent."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.6)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "微软雅黑"
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn("w:rFonts"), {})
    rFonts.set(qn("w:eastAsia"), "微软雅黑")
    rPr.insert(0, rFonts)
    return p


def _add_section_heading(doc, text):
    """Add a major section heading with bottom border."""
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)
    _add_heading_run(h, text, size=13, bold=True, color=RGBColor(30, 60, 120))
    pPr = h._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single", qn("w:sz"): "4",
        qn("w:space"): "1", qn("w:color"): "1E3C78",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_sub_heading(doc, text):
    """Add a sub-section heading."""
    sh = doc.add_paragraph()
    sh.paragraph_format.space_before = Pt(8)
    sh.paragraph_format.space_after = Pt(3)
    _add_heading_run(sh, f"▸ {text}", size=11, bold=True, color=RGBColor(60, 90, 150))


def _render_macro_section(doc, section, bold_title=True):
    """Render a macro-news section with numbered bold titles and indented content.
    Works for both 债市要闻速览 and DM信用早报."""
    items = section["items"]
    idx = 0
    news_num = 1

    while idx < len(items):
        title_item = items[idx]
        title_text = title_item["title"].strip()
        content_text = title_item.get("content", "").strip()

        # Handle merged title+content (split at first newline)
        if not content_text and "\n" in title_text:
            parts = title_text.split("\n", 1)
            title_text = parts[0].strip()
            content_text = parts[1].strip()

        # Check if next item is really content continuation
        if not content_text and idx + 1 < len(items):
            next_item = items[idx + 1]
            next_title = next_item["title"].strip()
            if is_content_line(next_title, title_text):
                content_text = next_title + "\n" + next_item.get("content", "")
                idx += 2
            else:
                idx += 1
        else:
            idx += 1

        # ── Numbered bold title ──
        tp = doc.add_paragraph()
        tp.paragraph_format.space_before = Pt(10)
        tp.paragraph_format.space_after = Pt(2)

        num_run = tp.add_run(f"{news_num}. ")
        num_run.font.size = Pt(10.5)
        num_run.font.bold = bold_title
        num_run.font.name = "微软雅黑"
        if bold_title:
            num_run.font.color.rgb = RGBColor(30, 60, 120)

        title_run = tp.add_run(title_text)
        title_run.font.size = Pt(10.5)
        title_run.font.bold = bold_title
        title_run.font.name = "微软雅黑"

        # ── Content ──
        if content_text:
            cp = doc.add_paragraph()
            cp.paragraph_format.space_after = Pt(6)
            cp.paragraph_format.first_line_indent = Cm(0.6)
            cr = cp.add_run(content_text)
            cr.font.size = Pt(10)
            cr.font.name = "微软雅黑"

        news_num += 1


def _render_simple_items(doc, items, bold_title=True):
    """Render items as simple bold-title + content paragraphs."""
    for item in items:
        title = item["title"].strip()
        content = item.get("content", "").strip()

        # If content is empty and title has newlines, split
        if not content and "\n" in title:
            parts = title.split("\n", 1)
            title = parts[0].strip()
            content = parts[1].strip()

        # Title
        tp = doc.add_paragraph()
        tp.paragraph_format.space_before = Pt(6)
        tp.paragraph_format.space_after = Pt(1)
        tr = tp.add_run(title)
        tr.font.size = Pt(10)
        tr.font.bold = bold_title
        tr.font.name = "微软雅黑"

        # Content
        if content:
            _add_body_paragraph(doc, content, indent=True, size=10)


def _render_bullet_items(doc, items):
    """Render items as simple bullet-point lines."""
    for item in items:
        title = item["title"].strip()
        content = item.get("content", "").strip()

        text = title
        if content:
            text += "：" + content

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.8)
        run = p.add_run(f"• {text}")
        run.font.size = Pt(9.5)
        run.font.name = "微软雅黑"


def build_docx(article):
    doc = Document()
    article_type = article.get("type", "yaowen")

    # ── Base style ──
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.35
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn("w:rFonts"), {})
    rFonts.set(qn("w:eastAsia"), "微软雅黑")
    rPr.insert(0, rFonts)

    # ── Title ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(20)
    title.paragraph_format.space_after = Pt(4)
    _add_heading_run(title, article["title"], size=16, bold=True)

    # ── Date & Tags ──
    if article["date"]:
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.paragraph_format.space_after = Pt(2)
        run = meta.add_run(f"发布日期：{article['date']}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)
        run.font.name = "微软雅黑"

    if article["tags"]:
        meta2 = doc.add_paragraph()
        meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta2.paragraph_format.space_after = Pt(6)
        run = meta2.add_run(f"标签：{article['tags']}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 149, 237)
        run.font.name = "微软雅黑"

    # ── Divider ──
    div = doc.add_paragraph()
    div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    div.paragraph_format.space_after = Pt(12)
    run = div.add_run("━" * 40)
    run.font.size = Pt(6)
    run.font.color.rgb = RGBColor(200, 200, 200)

    # ── Sections ──
    # Sections that should use numbered bold-title rendering (macro-style)
    MACRO_SECTIONS = ["宏观要闻", "【宏观要闻】", "公司新闻"]
    # Sections that should use simple bold-title rendering
    SIMPLE_SECTIONS = ["公司新闻", "机构观点", "【信用债】", "【重要资讯】",
                       "金融机构信息汇总"]
    # Sections that should use bullet rendering
    BULLET_SECTIONS = ["取消发行/终止审核", "财报业绩/企业拿地",
                       "评级动态", "人事动态", "经济数据"]

    last_heading = None  # Track for sub-section dedup

    for sec in article["sections"]:
        heading = sec["heading"]
        sub = sec.get("sub_heading")
        items = sec.get("items", [])

        if not items:
            continue

        # Section heading — skip if same as previous (sub-section continuation)
        if heading != last_heading:
            _add_section_heading(doc, heading)
            last_heading = heading

        # Sub-section heading (zaobao)
        if sub:
            _add_sub_heading(doc, sub)

        # Determine rendering style
        is_macro = any(m in heading for m in MACRO_SECTIONS)
        is_bullet = any(b in heading for b in BULLET_SECTIONS)
        is_simple = any(s in heading for s in SIMPLE_SECTIONS) or is_bullet

        if is_macro:
            bold = heading != "公司新闻"
            _render_macro_section(doc, sec, bold_title=bold)
        elif is_bullet:
            _render_bullet_items(doc, items)
        elif is_simple:
            bold = heading != "机构观点"
            _render_simple_items(doc, items, bold_title=bold)
        else:
            # Default: simple rendering
            _render_simple_items(doc, items)

    # ── Footer ──
    doc.add_paragraph()  # spacer
    div2 = doc.add_paragraph()
    div2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = div2.add_run("━" * 40)
    run.font.size = Pt(6)
    run.font.color.rgb = RGBColor(200, 200, 200)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(6)
    footer_text = "更多债市要闻详情，请查看DM终端舆情板块"
    run = footer.add_run(footer_text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.font.name = "微软雅黑"
    run.font.italic = True

    # ── Page setup ──
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    return doc


# ── Main entry ──────────────────────────────────────────────────────
def generate_docx(input_file, output_file=None):
    """Generate formatted DOCX from article text file.

    Args:
        input_file: path to the article text file
        output_file: path for output .docx (auto-generated if None)

    Returns:
        path to generated .docx file, or None on failure
    """
    lines, article_type = read_and_clean(input_file)
    article = parse_sections(lines, article_type)

    if not article["title"]:
        print("[ERROR] Could not extract article title")
        return None

    print(f"Type: {article_type}")
    print(f"Title: {article['title']}")
    print(f"Date: {article['date']}")
    print(f"Sections: {len(article['sections'])}")
    for sec in article["sections"]:
        print(f"  {sec['heading']}: {len(sec['items'])} items"
              + (f" [sub: {sec.get('sub_heading')}]" if sec.get('sub_heading') else ""))

    doc = build_docx(article)

    if output_file is None:
        output_file = os.path.join(OUT, "债市要闻速览.docx")

    doc.save(output_file)
    print(f"[OK] Saved: {output_file}")
    return output_file


if __name__ == "__main__":
    import argparse
    ap = argparse.ArgumentParser(description="Generate formatted DOCX from DM article")
    ap.add_argument("input", nargs="?", help="Input article text file")
    ap.add_argument("--output", "-o", help="Output DOCX file path")
    args = ap.parse_args()

    infile = args.input or INPUT
    generate_docx(infile, args.output)
